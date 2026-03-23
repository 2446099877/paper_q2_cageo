from __future__ import annotations

import argparse
import re
import shutil
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
PACKET_DIR = ROOT / "paper" / "submission_ready" / "cageo"
DOCS_DIR = ROOT / "docs" / "submission_packets"
REVIEWER_PACKET = ROOT / "paper" / "submission_ready" / "cageo_reviewer_code_packet.zip"
BUNDLE_DIR = ROOT / "paper" / "submission_ready" / "cageo_editorial_manager_bundle"
BUNDLE_ZIP = ROOT / "paper" / "submission_ready" / "cageo_editorial_manager_bundle.zip"


TEXT_DOCS = {
    "03_cover_letter": DOCS_DIR / "cageo_cover_letter_draft.md",
    "04_authorship_statement": DOCS_DIR / "cageo_authorship_statement.md",
    "05_competing_interest_statement": DOCS_DIR / "cageo_competing_interest_statement_2026-03-23.md",
    "06_generative_ai_declaration": DOCS_DIR / "cageo_generative_ai_declaration_draft_2026-03-23.md",
    "07_data_availability_statement": DOCS_DIR / "cageo_data_availability_statement_2026-03-23.md",
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Prepare an Editorial Manager upload bundle for the C&G submission."
    )
    parser.add_argument(
        "--output-dir",
        default=str(BUNDLE_DIR),
        help="Directory where the upload bundle will be written.",
    )
    return parser.parse_args()


def reset_dir(path: Path) -> None:
    if path.exists():
        shutil.rmtree(path)
    path.mkdir(parents=True, exist_ok=True)


def normalize_text(line: str) -> str:
    line = re.sub(r"`([^`]+)`", r"\1", line)
    line = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", line)
    return line.strip()


def markdown_lines(path: Path) -> list[str]:
    lines: list[str] = []
    for raw in path.read_text(encoding="utf-8").splitlines():
        stripped = raw.strip()
        if not stripped or stripped == "```":
            lines.append("")
            continue
        if stripped.startswith("更新时间："):
            continue
        lines.append(normalize_text(stripped))
    return lines


def write_text_and_docx(prefix: str, path: Path, output_dir: Path) -> None:
    lines = markdown_lines(path)
    txt_path = output_dir / f"{prefix}.txt"
    txt_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")

    document = Document()
    for line in lines:
        if not line:
            document.add_paragraph("")
            continue
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        else:
            document.add_paragraph(line)
    document.save(output_dir / f"{prefix}.docx")


def write_source_zip(output_dir: Path) -> None:
    source_zip = output_dir / "08_source_files.zip"
    include_files = [
        PACKET_DIR / "manuscript.tex",
        PACKET_DIR / "references.bib",
        PACKET_DIR / "cas-sc.cls",
        PACKET_DIR / "cas-common.sty",
        PACKET_DIR / "cas-model2-names.bst",
    ]
    with zipfile.ZipFile(source_zip, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in include_files:
            archive.write(path, arcname=path.name)
        for path in sorted((PACKET_DIR / "figures").rglob("*")):
            if path.is_file():
                archive.write(path, arcname=path.relative_to(PACKET_DIR).as_posix())
        for path in sorted((PACKET_DIR / "thumbnails").rglob("*")):
            if path.is_file():
                archive.write(path, arcname=path.relative_to(PACKET_DIR).as_posix())


def write_bundle_readme(output_dir: Path) -> None:
    text = """# C&G Editorial Manager Upload Bundle

This directory contains the upload-ready files for the current `Computers & Geosciences` submission.

## Recommended upload mapping

- `01_manuscript.pdf`: article PDF
- `02_highlights.txt`: highlights file
- `03_cover_letter.docx` or `03_cover_letter.txt`: cover letter
- `04_authorship_statement.docx` or `04_authorship_statement.txt`: authorship statement
- `05_competing_interest_statement.docx` or `05_competing_interest_statement.txt`: competing interest statement
- `06_generative_ai_declaration.docx` or `06_generative_ai_declaration.txt`: generative AI declaration
- `07_data_availability_statement.docx` or `07_data_availability_statement.txt`: data availability statement
- `08_source_files.zip`: manuscript source files for production
- `09_reviewer_code_packet.zip`: reviewer-safe supplementary material

## Notes

- Text and Word versions are both provided for convenience.
- `08_source_files.zip` includes the LaTeX manuscript, bibliography, class/style files, thumbnail assets, and figures.
- The reviewer code packet is optional but recommended for transparent review.
"""
    (output_dir / "README.md").write_text(text, encoding="utf-8")


def write_bundle_zip(output_dir: Path) -> None:
    if BUNDLE_ZIP.exists():
        BUNDLE_ZIP.unlink()
    with zipfile.ZipFile(BUNDLE_ZIP, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in sorted(output_dir.rglob("*")):
            if path.is_file():
                archive.write(path, arcname=path.relative_to(output_dir).as_posix())


def main() -> None:
    args = parse_args()
    output_dir = Path(args.output_dir).resolve()
    reset_dir(output_dir)

    shutil.copy2(PACKET_DIR / "manuscript.pdf", output_dir / "01_manuscript.pdf")
    shutil.copy2(PACKET_DIR / "highlights.txt", output_dir / "02_highlights.txt")

    for prefix, source in TEXT_DOCS.items():
        write_text_and_docx(prefix, source, output_dir)

    write_source_zip(output_dir)
    shutil.copy2(REVIEWER_PACKET, output_dir / "09_reviewer_code_packet.zip")
    write_bundle_readme(output_dir)
    write_bundle_zip(output_dir)
    print(f"prepared Editorial Manager bundle under {output_dir}")


if __name__ == "__main__":
    main()
